"""
generate_alignment_report.py
=============================
Generates valuesims_alignment_report.docx - a complete technical alignment
document showing every source file verbatim, divided by stream.

Run:
    python generate_alignment_report.py
"""

from __future__ import annotations

import os
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

# ── Configuration ─────────────────────────────────────────────────────────────

ROOT = Path(__file__).parent
OUTPUT = ROOT / "valuesims_alignment_report.docx"

EXCLUDE_PATTERNS = {
    "__pycache__",
    ".git",
    ".venv",
    "venv",
    "node_modules",
    ".pytest_cache",
    "ARTIFACTS",
    "dist",
    "build",
}
EXCLUDE_SUFFIXES = {".pyc", ".docx", ".xlsx", ".drawio", ".bkp"}
EXCLUDE_FILES = {"generate_alignment_report.py", "generate_project_docs.py"}

CODE_FONT = "Courier New"
CODE_SIZE = Pt(8)
CODE_BG_HEX = "F2F2F2"


# Stream classification by path prefix/name
# Returns (stream_label, badge, section_key)
def classify_file(rel: str) -> tuple[str, str, str]:
    parts = Path(rel).parts
    top = parts[0] if parts else ""
    sub = parts[1] if len(parts) > 1 else ""

    if top == "ui":
        return "UI / Infra (Tested)", "✅ TESTED", "tested"
    if top == "backend":
        name = Path(rel).name
        if name == "app.py" or name == "config.py":
            return "UI / Infra (Tested)", "✅ TESTED", "tested"
        if sub in (
            "infra",
            "ingestion",
            "retrieval",
            "state",
            "conversation",
        ) or name in ("app.py", "config.py"):
            return "UI / Infra (Tested)", "✅ TESTED", "tested"
        if sub in ("api",):
            return "UI / Infra (Tested)", "✅ TESTED", "tested"
        if sub in ("workflow", "ai", "tools") or name == "entry_handler.py":
            return "Agent (Untested)", "⚠️ UNTESTED", "agent"
        if sub == "entry":
            return "Agent (Untested)", "⚠️ UNTESTED", "agent"
    if top == "tests":
        return "Test", "🧪 TEST", "test"
    if top in ("docs",):
        return "Config / Docs", "📄 CONFIG", "config"
    return "Config / CI-CD", "⚙️ CONFIG", "config"


# ── First-test priority for agent node files ──────────────────────────────────

PRIORITY = {
    "backend/workflow/unified_incident_graph.py": "🔴 CRITICAL",
    "backend/ai/escalation_controller.py": "🔴 CRITICAL",
    "backend/workflow/nodes/operational_node.py": "🔴 CRITICAL",
    "backend/workflow/nodes/operational_reflection_node.py": "🔴 CRITICAL",
    "backend/workflow/nodes/operational_escalation_node.py": "🔴 CRITICAL",
    "backend/workflow/nodes/intent_classification_node.py": "🟡 NEXT",
    "backend/workflow/nodes/router_node.py": "🟡 NEXT",
    "backend/workflow/nodes/response_formatter_node.py": "🟡 NEXT",
    "backend/workflow/nodes/intent_reflection_node.py": "🟡 NEXT",
    "backend/workflow/nodes/similarity_node.py": "🟢 LATER",
    "backend/workflow/nodes/similarity_reflection_node.py": "🟢 LATER",
    "backend/workflow/nodes/strategy_node.py": "🟢 LATER",
    "backend/workflow/nodes/strategy_reflection_node.py": "🟢 LATER",
    "backend/workflow/nodes/strategy_escalation_node.py": "🟢 LATER",
    "backend/workflow/nodes/kpi_node.py": "🟢 LATER",
    "backend/workflow/nodes/kpi_reflection_node.py": "🟢 LATER",
    "backend/workflow/nodes/start_node.py": "🟢 LATER",
    "backend/workflow/nodes/context_node.py": "🟡 NEXT",
    "backend/workflow/nodes/end_node.py": "🟢 LATER",
    "backend/ai/model_policy.py": "🟡 NEXT",
    "backend/ai/model_strategy.py": "🟡 NEXT",
    "backend/tools/kpi_tool.py": "🟢 LATER",
    "backend/entry/entry_handler.py": "🔴 CRITICAL",
    "backend/workflow/models.py": "🔴 CRITICAL",
}


# ── File collection ───────────────────────────────────────────────────────────


def should_skip(path: Path) -> bool:
    for part in path.parts:
        if part in EXCLUDE_PATTERNS:
            return True
    if path.suffix in EXCLUDE_SUFFIXES:
        return True
    if path.name in EXCLUDE_FILES:
        return True
    return False


def collect(root: Path) -> list[Path]:
    return sorted([p for p in root.rglob("*") if p.is_file() and not should_skip(p)])


def rel(p: Path) -> str:
    return p.relative_to(ROOT).as_posix()


def read_verbatim(p: Path) -> tuple[str, bool]:
    """Returns (content, is_empty)"""
    try:
        text = p.read_text(encoding="utf-8", errors="replace")
        if not text.strip():
            return "[EMPTY FILE]", True
        return text, False
    except Exception as exc:
        return f"[READ ERROR: {exc}]", False


# ── DocX helpers ──────────────────────────────────────────────────────────────


def set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def add_left_border(paragraph, hex_color: str, width_pt: int = 12):
    """Add a thick left border (colored bar) to a paragraph."""
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), str(width_pt * 8))
    left.set(qn("w:space"), "4")
    left.set(qn("w:color"), hex_color)
    pBdr.append(left)
    pPr.append(pBdr)


def add_toc(doc: Document):
    para = doc.add_paragraph()
    run = para.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-3" \\h \\z \\u'
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_sep)
    run._r.append(fld_end)


def add_code_block(doc: Document, code: str, border_color: str):
    """Add a verbatim code block with gray background and colored left border."""
    lines = code.splitlines()
    if not lines:
        lines = [""]

    # Use a single paragraph approach with line-by-line newlines to avoid
    # massive paragraph count — group into chunks of 200 lines.
    CHUNK = 200
    for chunk_start in range(0, len(lines), CHUNK):
        chunk_lines = lines[chunk_start : chunk_start + CHUNK]
        chunk_text = "\n".join(chunk_lines)

        para = doc.add_paragraph()
        add_left_border(para, border_color)

        # Gray paragraph background via paragraph shading
        pPr = para._p.get_or_add_pPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), CODE_BG_HEX)
        pPr.append(shd)

        run = para.add_run(chunk_text)
        run.font.name = CODE_FONT
        run.font.size = CODE_SIZE
        run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)


def add_notes_field(doc: Document, label: str):
    """Add a blank labelled field for handwritten review notes."""
    para = doc.add_paragraph()
    run = para.add_run(f"{label} ")
    run.bold = True
    run.font.color.rgb = RGBColor(0xCC, 0x44, 0x00)

    # Dotted lines for handwriting
    dots = doc.add_paragraph("_" * 110)
    dots.runs[0].font.color.rgb = RGBColor(0xBB, 0xBB, 0xBB)
    dots.runs[0].font.size = Pt(10)
    doc.add_paragraph()


def file_description(rel_path: str, content: str) -> str:
    """
    Return a short (2-3 sentence) role description by inspecting actual content.
    """
    lines = content.splitlines()
    # Gather docstring or first meaningful comment
    for line in lines[:30]:
        stripped = line.strip()
        if stripped.startswith('"""') or stripped.startswith("'''"):
            return stripped.strip("\"'").strip() or f"Source file: {rel_path}"
        if stripped.startswith("#") and len(stripped) > 5:
            return stripped.lstrip("#").strip()

    # Fallback: describe by imports + class/function names
    classes = [l.strip() for l in lines if l.strip().startswith("class ")][:3]
    funcs = [l.strip() for l in lines if l.strip().startswith("def ")][:3]
    parts = []
    if classes:
        parts.append(
            "Classes: "
            + ", ".join(c.split("(")[0].replace("class ", "") for c in classes)
        )
    if funcs:
        parts.append(
            "Functions: "
            + ", ".join(f.split("(")[0].replace("def ", "") for f in funcs)
        )
    if parts:
        return ". ".join(parts) + "."
    return f"Source file ({Path(rel_path).suffix or 'no-ext'})."


# ── Section rendering ─────────────────────────────────────────────────────────

BORDER_TESTED = "2E7D32"  # dark green
BORDER_AGENT = "E65100"  # orange
BORDER_TEST = "1565C0"  # blue
BORDER_CONFIG = "6A1B9A"  # purple
BADGE_BG = {
    "tested": "E8F5E9",
    "agent": "FFF3E0",
    "test": "E3F2FD",
    "config": "F3E5F5",
}


def border_for_section(section_key: str) -> str:
    return {
        "tested": BORDER_TESTED,
        "agent": BORDER_AGENT,
        "test": BORDER_TEST,
        "config": BORDER_CONFIG,
    }.get(section_key, "888888")


def render_file_entry(
    doc: Document,
    file_path: Path,
    rel_path: str,
    content: str,
    stream_label: str,
    badge: str,
    section_key: str,
    with_alignment_notes: bool = False,
    with_test_status: bool = False,
):
    border = border_for_section(section_key)

    # Heading (level 3)
    h = doc.add_heading(rel_path, level=3)
    # Colour heading to match stream
    for run in h.runs:
        run.font.color.rgb = RGBColor.from_string(border)

    # Badge + meta table
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"

    def add_row(label: str, value: str, bg: str = "FFFFFF"):
        row = table.add_row()
        set_cell_bg(row.cells[0], "E0E0E0")
        set_cell_bg(row.cells[1], bg)
        row.cells[0].text = label
        row.cells[0].paragraphs[0].runs[0].bold = True
        row.cells[0].paragraphs[0].runs[0].font.size = Pt(9)
        row.cells[1].text = value
        row.cells[1].paragraphs[0].runs[0].font.size = Pt(9)

    add_row("Stream", f"{badge}  {stream_label}", BADGE_BG.get(section_key, "FFFFFF"))
    add_row("Path", rel_path)
    add_row("Extension", Path(rel_path).suffix or "(none)")
    desc = file_description(rel_path, content)
    add_row("Role", desc)

    doc.add_paragraph()

    # Code block
    add_code_block(doc, content, border)
    doc.add_paragraph()

    if with_alignment_notes:
        add_notes_field(doc, "ALIGNMENT NOTES:")

    if with_test_status:
        add_notes_field(doc, "TEST STATUS:")


# ── Main document builder ─────────────────────────────────────────────────────


def build(root: Path, output: Path):
    files = collect(root)

    # Classify
    tested_files, agent_files, test_files, config_files = [], [], [], []
    empty_count = 0

    for fp in files:
        r = rel(fp)
        _, _, section_key = classify_file(r)
        if section_key == "tested":
            tested_files.append(fp)
        elif section_key == "agent":
            agent_files.append(fp)
        elif section_key == "test":
            test_files.append(fp)
        else:
            config_files.append(fp)

    doc = Document()

    # ── Page margins ──────────────────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin = Inches(0.9)
        section.bottom_margin = Inches(0.9)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # ──────────────────────────────────────────────────────────────────────────
    # COVER PAGE
    # ──────────────────────────────────────────────────────────────────────────
    title = doc.add_heading("valuesims-decision-support", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = subtitle.add_run("Codebase Alignment Report — Technical Handover\n")
    r1.bold = True
    r1.font.size = Pt(14)
    subtitle.add_run(f"Generated: {date.today().isoformat()}\n").font.size = Pt(11)
    subtitle.add_run(
        "Purpose: Align Stream A (ChatGPT UI/Infra — TESTED) "
        "with Stream B (Claude Agent Layer — UNTESTED)\n"
        "Immediate Goal: Verify operational node test readiness\n"
    ).font.size = Pt(10)

    doc.add_paragraph()

    # Status table
    doc.add_heading("Stream Status Summary", level=2)
    tbl = doc.add_table(rows=4, cols=3)
    tbl.style = "Table Grid"

    headers = ["Stream", "Files", "Status"]
    rows_data = [
        headers,
        [
            "UI / Infra (ChatGPT — Stream A)",
            str(len(tested_files)),
            "✅ TESTED — DO NOT BREAK",
        ],
        [
            "Agent Layer (Claude — Stream B)",
            str(len(agent_files)),
            "⚠️ UNTESTED — Alignment in progress",
        ],
        ["Tests", str(len(test_files)), "🧪 TEST FILES — Run before first live test"],
    ]
    bg_colors = ["D0D0D0", "E8F5E9", "FFF3E0", "E3F2FD"]
    for i, (row_data, bg) in enumerate(zip(rows_data, bg_colors)):
        row = tbl.rows[i]
        for j, (cell, val) in enumerate(zip(row.cells, row_data)):
            set_cell_bg(cell, bg)
            cell.text = val
            cell.paragraphs[0].runs[0].bold = i == 0
            cell.paragraphs[0].runs[0].font.size = Pt(10)

    doc.add_paragraph()

    priority_note = doc.add_paragraph()
    priority_note.add_run("OPERATIONAL NODE TEST CHECKLIST  ").bold = True
    chk_items = [
        "OperationalNode correctly calls LoggedLanguageModelClient.complete_json()",
        "OperationalReflectionNode returns parseable ReflectionResult with needs_escalation bool",
        "EscalationController.should_escalate_operational() reads operational_reflection from state",
        "UnifiedIncidentGraph conditional edge routes ESCALATE / CONTINUE correctly",
        "All node imports resolve — no circular imports, no missing modules",
    ]
    for item in chk_items:
        p = doc.add_paragraph(f"  ☐ {item}", style="List Bullet")
        p.runs[0].font.size = Pt(10)

    doc.add_page_break()

    # ──────────────────────────────────────────────────────────────────────────
    # TABLE OF CONTENTS
    # ──────────────────────────────────────────────────────────────────────────
    doc.add_heading("Table of Contents", level=1)
    toc_note = doc.add_paragraph(
        "Right-click and select 'Update Field' in Microsoft Word to populate the TOC."
    )
    toc_note.italic = True
    add_toc(doc)
    doc.add_page_break()

    # ──────────────────────────────────────────────────────────────────────────
    # SECTION 1 — TESTED FILES
    # ──────────────────────────────────────────────────────────────────────────
    doc.add_heading("Section 1 — TESTED FILES  (UI / Infra — Stream A)", level=1)
    intro = doc.add_paragraph()
    intro.add_run("Stream: ").bold = True
    intro.add_run("UI / ChatGPT-built  |  ")
    intro.add_run("Status: ").bold = True
    intro.add_run("✅ TESTED AND WORKING — do not modify without regression testing.")
    intro.add_run(
        "\nCovers: ui/, backend/infra/, backend/ingestion/, backend/retrieval/, "
        "backend/state/, backend/conversation/, backend/api/, backend/app.py, backend/config.py"
    ).font.size = Pt(9)
    doc.add_paragraph()

    tested_empty = 0
    for fp in tested_files:
        r_path = rel(fp)
        content, is_empty = read_verbatim(fp)
        if is_empty:
            tested_empty += 1
        stream_label, badge, section_key = classify_file(r_path)
        render_file_entry(
            doc,
            fp,
            r_path,
            content,
            stream_label,
            badge,
            section_key,
            with_alignment_notes=False,
            with_test_status=False,
        )

    doc.add_page_break()

    # ──────────────────────────────────────────────────────────────────────────
    # SECTION 2 — AGENT FILES (UNTESTED)
    # ──────────────────────────────────────────────────────────────────────────
    doc.add_heading("Section 2 — AGENT FILES  (Claude Agent Layer — Stream B)", level=1)
    intro2 = doc.add_paragraph()
    intro2.add_run("Stream: ").bold = True
    intro2.add_run("Claude Agent Layer  |  ")
    intro2.add_run("Status: ").bold = True
    intro2.add_run("⚠️ UNTESTED — Review each file before first operational node test.")
    intro2.add_run(
        "\nCovers: backend/workflow/, backend/ai/, backend/tools/, backend/entry/"
    ).font.size = Pt(9)
    doc.add_paragraph()

    agent_empty = 0
    for fp in agent_files:
        r_path = rel(fp)
        content, is_empty = read_verbatim(fp)
        if is_empty:
            agent_empty += 1
        stream_label, badge, section_key = classify_file(r_path)
        # Flag __init__ files
        is_node_file = "nodes/" in r_path or r_path in PRIORITY
        render_file_entry(
            doc,
            fp,
            r_path,
            content,
            stream_label,
            badge,
            section_key,
            with_alignment_notes=is_node_file,
            with_test_status=False,
        )

    doc.add_page_break()

    # ──────────────────────────────────────────────────────────────────────────
    # SECTION 3 — TEST FILES
    # ──────────────────────────────────────────────────────────────────────────
    doc.add_heading("Section 3 — TEST FILES", level=1)
    doc.add_paragraph(
        "🧪 Run all tests before the first live operational node test. "
        "Tests cover escalation controller, model policy, and operational adaptive flow."
    )
    doc.add_paragraph()

    test_empty = 0
    for fp in test_files:
        r_path = rel(fp)
        content, is_empty = read_verbatim(fp)
        if is_empty:
            test_empty += 1
        stream_label, badge, section_key = classify_file(r_path)
        render_file_entry(
            doc,
            fp,
            r_path,
            content,
            stream_label,
            badge,
            section_key,
            with_alignment_notes=False,
            with_test_status=True,
        )

    doc.add_page_break()

    # ──────────────────────────────────────────────────────────────────────────
    # SECTION 4 — CONFIG & DOCS
    # ──────────────────────────────────────────────────────────────────────────
    doc.add_heading("Section 4 — Config, Docs & CI/CD", level=1)
    doc.add_paragraph("Root-level config, documentation, and GitHub workflows.")
    doc.add_paragraph()

    config_empty = 0
    for fp in config_files:
        r_path = rel(fp)
        content, is_empty = read_verbatim(fp)
        if is_empty:
            config_empty += 1
        stream_label, badge, section_key = classify_file(r_path)
        render_file_entry(
            doc,
            fp,
            r_path,
            content,
            stream_label,
            badge,
            section_key,
        )

    doc.add_page_break()

    # ──────────────────────────────────────────────────────────────────────────
    # SECTION 5 — ALIGNMENT SUMMARY TABLE
    # ──────────────────────────────────────────────────────────────────────────
    doc.add_heading("Section 5 — Alignment Summary Table (Agent Files)", level=1)
    doc.add_paragraph(
        "One row per agent file. Review expected behaviour against actual code."
    )
    doc.add_paragraph()

    agent_table = doc.add_table(rows=1, cols=5)
    agent_table.style = "Table Grid"
    headers5 = ["File", "Exists", "Has Code", "Imports Resolve?", "First Test Priority"]
    hdr_row = agent_table.rows[0]
    for cell, hdr in zip(hdr_row.cells, headers5):
        set_cell_bg(cell, "263238")
        cell.text = hdr
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        cell.paragraphs[0].runs[0].font.size = Pt(9)

    # All agent files + priority files not already in list
    all_agent_rels = {rel(fp): fp for fp in agent_files}
    priority_rels = set(PRIORITY.keys())
    combined = sorted(set(list(all_agent_rels.keys()) + list(priority_rels)))

    priority_bg = {
        "🔴 CRITICAL": "FFEBEE",
        "🟡 NEXT": "FFFDE7",
        "🟢 LATER": "F1F8E9",
        "—": "FAFAFA",
    }

    for r_path in combined:
        fp = all_agent_rels.get(r_path)
        exists = "✅ Yes" if fp and fp.exists() else "❌ No"
        has_code = "—"
        imports_ok = "—"
        if fp and fp.exists():
            content, is_empty = read_verbatim(fp)
            has_code = "[EMPTY]" if is_empty else "✅ Yes"
            # Very basic import check: look for any import lines
            import_lines = [
                l
                for l in content.splitlines()
                if l.strip().startswith(("import ", "from "))
            ]
            imports_ok = "✅ Present" if import_lines else "⚠️ No imports"

        prio = PRIORITY.get(r_path, "—")
        bg = priority_bg.get(prio, "FAFAFA")

        data_row = agent_table.add_row()
        values = [Path(r_path).name, exists, has_code, imports_ok, prio]
        for cell, val in zip(data_row.cells, values):
            set_cell_bg(cell, bg)
            cell.text = val
            cell.paragraphs[0].runs[0].font.size = Pt(8)

    doc.add_paragraph()

    # Critical findings callout
    doc.add_heading("Critical Alignment Findings", level=2)
    findings = [
        (
            "OperationalNode.run(model_name=...)",
            "EXISTS — calls self._llm_client.complete_text() with model_name override. "
            "OperationalEscalationNode calls run(model_name=...) with model_name from ModelPolicy. ✅",
        ),
        (
            "OperationalReflectionNode → ReflectionResult",
            "EXISTS — returns OperationalReflectionOutput containing ReflectionResult with "
            "needs_escalation bool derived from completeness_score < 0.65 OR hallucination_risk=='HIGH' "
            "OR schema_invalid OR issues present. ✅",
        ),
        (
            "EscalationController reads state",
            "EXISTS — reads state['operational_reflection'] for needs_escalation bool, "
            "guards against double-escalation via state['operational_escalated']. ✅",
        ),
        (
            "UnifiedIncidentGraph conditional edge",
            "EXISTS — _route_operational_escalation() calls escalation_controller.should_escalate_operational(). "
            "Returns 'ESCALATE'→operational_escalation_node or 'CONTINUE'→response_formatter_node. ✅",
        ),
        (
            "operational_escalation_node → operational_reflection_node edge",
            "EXISTS — graph.add_edge('operational_escalation_node', 'operational_reflection_node'). "
            "This creates a reflection loop. Guard via operational_escalated=True in state. ✅",
        ),
        (
            "Import resolution risk",
            "backend/ingestion/case_ingestion.py imports IncidentFactory, LegacyCaseModel, "
            "IncidentStateAdapter from backend.state.incident_state — these all exist. "
            "No circular imports detected from static inspection. ✅",
        ),
        (
            "KnowledgeIngestionService hardcodes embedding length=3072",
            "⚠️ RISK: Line checks len(embedding) != 3072. "
            "If your Azure OpenAI embedding deployment uses a different dim (e.g. 1536), "
            "this will raise ValueError on every knowledge upload. Verify deployment dimension.",
        ),
        (
            "EmbeddingClient — no default model name",
            "⚠️ RISK: Falls back to os.environ vars. "
            "Settings has AZURE_OPENAI_CHAT_DEPLOYMENT but EmbeddingClient does NOT read from Settings — "
            "it reads directly from os.environ. Ensure AZURE_OPENAI_EMBEDDING_DEPLOYMENT env var is set.",
        ),
    ]

    findings_tbl = doc.add_table(rows=1, cols=2)
    findings_tbl.style = "Table Grid"
    fh = findings_tbl.rows[0]
    for cell, hdr in zip(fh.cells, ["Finding", "Assessment"]):
        set_cell_bg(cell, "37474F")
        cell.text = hdr
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        cell.paragraphs[0].runs[0].font.size = Pt(9)

    for finding, assessment in findings:
        row = findings_tbl.add_row()
        bg = "FFF8E1" if "⚠️" in assessment else "F1F8E9"
        set_cell_bg(row.cells[0], "ECEFF1")
        set_cell_bg(row.cells[1], bg)
        row.cells[0].text = finding
        row.cells[0].paragraphs[0].runs[0].font.size = Pt(8)
        row.cells[0].paragraphs[0].runs[0].bold = True
        row.cells[1].text = assessment
        row.cells[1].paragraphs[0].runs[0].font.size = Pt(8)

    doc.add_page_break()

    # ──────────────────────────────────────────────────────────────────────────
    # APPENDIX — FLAT FILE LIST
    # ──────────────────────────────────────────────────────────────────────────
    doc.add_heading("Appendix — All Documented File Paths", level=1)
    doc.add_paragraph("Flat list of all files included in this report.")
    doc.add_paragraph()

    all_documented = tested_files + agent_files + test_files + config_files
    for fp in sorted(all_documented, key=rel):
        r_path = rel(fp)
        _, badge, _ = classify_file(r_path)
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(f"{badge}  {r_path}")
        run.font.size = Pt(8)
        run.font.name = CODE_FONT

    # ── Save ─────────────────────────────────────────────────────────────────
    doc.save(str(output))

    total = len(all_documented)
    total_empty = tested_empty + agent_empty + test_empty + config_empty
    return {
        "total": total,
        "tested": len(tested_files),
        "agent": len(agent_files),
        "test": len(test_files),
        "config": len(config_files),
        "empty": total_empty,
    }


# ── Entry point ───────────────────────────────────────────────────────────────

if __name__ == "__main__":
    print(f"Root : {ROOT}")
    print(f"Output: {OUTPUT}")
    print("Building alignment report …\n")

    stats = build(ROOT, OUTPUT)

    print(f"✓  Saved: {OUTPUT}")
    size_kb = OUTPUT.stat().st_size / 1024
    print(f"   File size       : {size_kb:.1f} KB")
    print(f"\n   Files documented: {stats['total']}")
    print(f"   Tested stream   : {stats['tested']}")
    print(f"   Untested agent  : {stats['agent']}")
    print(f"   Test files      : {stats['test']}")
    print(f"   Config/docs     : {stats['config']}")
    print(f"   Empty files     : {stats['empty']}")
