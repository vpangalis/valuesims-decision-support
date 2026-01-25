"""Export repository files into a single Word (.docx) document.

- Includes a directory tree at the top.
- Then includes each file with its relative path and contents.
- Skips secrets and noisy/binary paths by default.

Usage (from backend/):
  python scripts/export_repo_to_docx.py --repo-root .. --out ../valuesims-decision-support.docx

You can also run from repo root:
  python backend/scripts/export_repo_to_docx.py --repo-root . --out valuesims-decision-support.docx

Notes:
- This script intentionally excludes .env and similar secret files.
- It also excludes .git/, .venv/, __pycache__/ and common binary/media extensions.
"""

from __future__ import annotations

import argparse
import os
from pathlib import Path


TEXT_EXTENSIONS = {
    ".py",
    ".md",
    ".txt",
    ".json",
    ".yml",
    ".yaml",
    ".toml",
    ".ini",
    ".cfg",
    ".html",
    ".css",
    ".js",
    ".ts",
    ".tsx",
    ".jsx",
    ".sh",
    ".ps1",
    ".bat",
    ".sql",
}

BINARY_EXTENSIONS = {
    ".png",
    ".jpg",
    ".jpeg",
    ".gif",
    ".bmp",
    ".ico",
    ".svg",
    ".pdf",
    ".zip",
    ".tar",
    ".gz",
    ".7z",
    ".exe",
    ".dll",
    ".so",
    ".dylib",
    ".pyd",
    ".pyo",
    ".pyc",
    ".pkl",
    ".db",
    ".sqlite",
    ".docx",
    ".xlsx",
    ".pptx",
}

EXCLUDE_DIR_NAMES = {
    ".git",
    ".github",
    ".venv",
    "venv",
    "__pycache__",
    "node_modules",
    ".pytest_cache",
    ".mypy_cache",
}

DOTENV_FILE_NAMES = {
    ".env",
    ".env.local",
    ".env.development",
    ".env.production",
    ".env.test",
}


def should_exclude_path(path: Path, *, include_env: bool) -> bool:
    parts = {p.lower() for p in path.parts}
    if any(p in EXCLUDE_DIR_NAMES for p in parts):
        return True

    if not include_env and path.name in DOTENV_FILE_NAMES:
        return True

    # Exclude typical key/cert artifacts
    lowered = path.name.lower()
    if lowered.endswith((".pem", ".key", ".pfx", ".p12")):
        return True

    return False


def is_text_file(path: Path) -> bool:
    suffix = path.suffix.lower()
    if suffix in BINARY_EXTENSIONS:
        return False
    if suffix in TEXT_EXTENSIONS:
        return True

    # Fallback: treat no-extension files as text (README, commit_commands, etc.)
    return suffix == ""


def build_tree(repo_root: Path, *, include_env: bool, max_entries: int = 2000) -> str:
    lines: list[str] = []
    count = 0

    def walk(dir_path: Path, prefix: str = "") -> None:
        nonlocal count
        if count >= max_entries:
            return

        entries = sorted(dir_path.iterdir(), key=lambda p: (p.is_file(), p.name.lower()))
        for idx, entry in enumerate(entries):
            if should_exclude_path(entry, include_env=include_env):
                continue
            if count >= max_entries:
                break

            connector = "└── " if idx == len(entries) - 1 else "├── "
            rel = entry.relative_to(repo_root).as_posix()
            lines.append(f"{prefix}{connector}{rel}")
            count += 1

            if entry.is_dir():
                extension = "    " if idx == len(entries) - 1 else "│   "
                walk(entry, prefix + extension)

    lines.append(repo_root.name)
    walk(repo_root)

    if count >= max_entries:
        lines.append("… (tree truncated)")

    return "\n".join(lines)


def iter_files(repo_root: Path, *, include_env: bool) -> list[Path]:
    files: list[Path] = []
    for root, dirnames, filenames in os.walk(repo_root):
        root_path = Path(root)

        # Mutate dirnames in-place to prune traversal
        dirnames[:] = [
            d
            for d in dirnames
            if d not in EXCLUDE_DIR_NAMES
            and not should_exclude_path(root_path / d, include_env=include_env)
        ]

        for filename in filenames:
            path = root_path / filename
            if should_exclude_path(path, include_env=include_env):
                continue
            if not path.is_file():
                continue
            files.append(path)

    return sorted(files, key=lambda p: p.relative_to(repo_root).as_posix().lower())


def main() -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument("--repo-root", type=str, required=True, help="Path to repo root")
    parser.add_argument("--out", type=str, required=True, help="Output .docx path")
    parser.add_argument(
        "--max-bytes-per-file",
        type=int,
        default=300_000,
        help="Skip/clip very large files (default: 300KB)",
    )
    parser.add_argument(
        "--include-env",
        action="store_true",
        help="Include .env files in the export (may contain secrets)",
    )
    args = parser.parse_args()

    repo_root = Path(args.repo_root).resolve()
    out_path = Path(args.out).resolve()

    if not repo_root.exists() or not repo_root.is_dir():
        raise ValueError(f"Invalid --repo-root: {repo_root}")
    if out_path.suffix.lower() != ".docx":
        raise ValueError("--out must end with .docx")

    try:
        from docx import Document  # type: ignore
        from docx.shared import Pt  # type: ignore
    except Exception as exc:  # pragma: no cover
        raise RuntimeError(
            "Missing dependency 'python-docx'. Install with: pip install python-docx"
        ) from exc

    document = Document()
    document.add_heading(f"Repository export: {repo_root.name}", level=0)

    document.add_heading("Directory tree", level=1)
    tree_text = build_tree(repo_root, include_env=args.include_env)
    p = document.add_paragraph()
    run = p.add_run(tree_text)
    run.font.name = "Courier New"
    run.font.size = Pt(9)

    document.add_page_break()
    document.add_heading("Files", level=1)

    for path in iter_files(repo_root, include_env=args.include_env):
        rel = path.relative_to(repo_root).as_posix()
        if not is_text_file(path):
            continue

        try:
            data = path.read_bytes()
        except Exception:
            continue

        if len(data) > args.max_bytes_per_file:
            # Clip large files rather than dumping megabytes into Word
            data = data[: args.max_bytes_per_file]
            clipped = True
        else:
            clipped = False

        try:
            text = data.decode("utf-8")
        except UnicodeDecodeError:
            # Fallback for odd encodings
            text = data.decode("utf-8", errors="replace")

        document.add_heading(rel, level=2)
        if clipped:
            document.add_paragraph(f"(clipped to {args.max_bytes_per_file} bytes)")

        p = document.add_paragraph()
        run = p.add_run(text)
        run.font.name = "Courier New"
        run.font.size = Pt(9)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(out_path)

    print(f"Wrote: {out_path}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
