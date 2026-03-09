"""
generate_project_docs.py
========================
Generates a comprehensive Word document (.docx) cataloging every file in the
valuesims-decision-support project.

Run:
    python generate_project_docs.py
"""

from __future__ import annotations

import os
import re
import textwrap
from datetime import date
from pathlib import Path
from typing import Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Inches

# ---------------------------------------------------------------------------
# Configuration
# ---------------------------------------------------------------------------

PROJECT_ROOT = Path(__file__).parent
OUTPUT_FILE = PROJECT_ROOT / "valuesims_decision_support_project_files.docx"

EXCLUDED_DIRS = {
    "__pycache__",
    ".git",
    ".venv",
    "venv",
    "node_modules",
    "dist",
    "build",
}
EXCLUDED_EXTENSIONS = {".pyc"}

SUMMARISABLE_EXTENSIONS = {
    ".py",
    ".ts",
    ".json",
    ".yaml",
    ".yml",
    ".md",
    ".env",
    ".ini",
    ".cfg",
    ".toml",
    ".js",
    ".css",
    ".html",
}

# Maximum characters to read per file for summary extraction
MAX_READ_CHARS = 6000

# ---------------------------------------------------------------------------
# Curated per-file descriptions
# (relative POSIX paths → description dict)
# These override auto-generated descriptions with richer domain knowledge.
# ---------------------------------------------------------------------------

CURATED: dict[str, dict] = {
    # ── Root ──────────────────────────────────────────────────────────────
    "pytest.ini": {
        "purpose": "Pytest configuration file for the test suite.",
        "summary": (
            "Defines pytest settings including test discovery paths, markers, "
            "and output verbosity. Controls how unit and integration tests are "
            "collected and run across the project."
        ),
    },
    # ── Backend root ──────────────────────────────────────────────────────
    "backend/__init__.py": {
        "purpose": "Package initialiser for the backend module.",
        "summary": "Marks the backend directory as a Python package.",
    },
    "backend/app.py": {
        "purpose": "FastAPI application factory and dependency-injection container.",
        "summary": (
            "Defines BackendContainer, which wires together all infrastructure "
            "clients (Blob Storage, Azure AI Search, OpenAI embeddings, LLM), "
            "ingestion services, workflow nodes, and the UnifiedIncidentGraph. "
            "Creates the FastAPI app instance, registers CORS middleware, and "
            "mounts the API router. Acts as the composition root for the entire "
            "backend."
        ),
    },
    "backend/config.py": {
        "purpose": "Centralised application configuration via Pydantic BaseSettings.",
        "summary": (
            "Declares a Settings class that reads all environment variables "
            "required by the platform: Azure Search endpoint and keys, Blob "
            "Storage connection strings, index names (case / evidence / "
            "knowledge), retrieval top-k values, OpenAI model names "
            "(gpt-4o-mini default, gpt-4o premium), and logging flags. "
            "A singleton `settings` instance is exported for import throughout "
            "the codebase."
        ),
    },
    # ── AI ────────────────────────────────────────────────────────────────
    "backend/ai/__init__.py": {
        "purpose": "Package initialiser for the ai sub-module.",
        "summary": "Marks the ai directory as a Python package.",
    },
    "backend/ai/escalation_controller.py": {
        "purpose": "Controls whether a LangGraph node should escalate to a premium model.",
        "summary": (
            "EscalationController exposes two methods: "
            "should_escalate_operational() and should_escalate_strategy(). "
            "Each inspects the LangGraph state dict for a reflection result that "
            "carries needs_escalation=True and guards against repeated escalation "
            "via the escalated flags in state. Used by the unified incident graph "
            "as conditional edge logic."
        ),
    },
    "backend/ai/model_policy.py": {
        "purpose": "Resolves which Azure OpenAI deployment to use for each graph node.",
        "summary": (
            "ModelPolicy wraps ModelStrategy and implements resolve_model(node_name, "
            "state). It returns the premium model deployment when a node has been "
            "escalated, and the default model otherwise. Supports 'operational', "
            "'strategy', and a fallback intent-classification model."
        ),
    },
    "backend/ai/model_strategy.py": {
        "purpose": "Value object holding Azure OpenAI deployment names from settings.",
        "summary": (
            "ModelStrategy reads five model deployment names from the Settings "
            "object: intent_default, operational_default, operational_premium, "
            "strategy_default, and strategy_premium. Provides a single source of "
            "truth for model routing across the LangGraph workflow."
        ),
    },
    # ── API ───────────────────────────────────────────────────────────────
    "backend/api/__init__.py": {
        "purpose": "Package initialiser for the api sub-module.",
        "summary": "Marks the api directory as a Python package.",
    },
    "backend/api/routes.py": {
        "purpose": "FastAPI router that exposes the three main HTTP entry points.",
        "summary": (
            "ApiRoutes defines POST /entry/case (case ingestion & updates), "
            "POST /entry/reasoning (AI reasoning invocation), and "
            "POST /entry/knowledge (binary file upload for knowledge ingestion). "
            "Each endpoint validates the intent field in the incoming envelope, "
            "normalises the action, and delegates to EntryHandler. Returns "
            "structured JSON responses or HTTP error codes."
        ),
    },
    # ── Conversation ──────────────────────────────────────────────────────
    "backend/conversation/conversation_handler.py": {
        "purpose": "Placeholder / handler for multi-turn conversation state management.",
        "summary": (
            "Provides scaffolding for handling conversational context across "
            "multiple AI reasoning turns. Intended to track session history and "
            "manage continuity of the D1–D8 problem-solving workflow across "
            "user interactions."
        ),
    },
    # ── Entry ─────────────────────────────────────────────────────────────
    "backend/entry/entry_handler.py": {
        "purpose": "Application-layer dispatcher that routes incoming envelopes to services.",
        "summary": (
            "EntryHandler receives an EntryEnvelope (intent + action + payload) "
            "and dispatches to the appropriate service: case ingestion / update / "
            "close, evidence upload, knowledge upload, or UnifiedIncidentGraph "
            "invocation for AI reasoning. Returns a structured "
            "EntryResponseEnvelope. Acts as the single application entry point "
            "called by all API routes."
        ),
    },
    # ── Infra ─────────────────────────────────────────────────────────────
    "backend/infra/blob_storage.py": {
        "purpose": "Azure Blob Storage client, CaseRepository, and CaseReadRepository.",
        "summary": (
            "BlobStorageClient wraps the Azure SDK ContainerClient for uploading, "
            "downloading, and deleting blobs. CaseRepository provides case-oriented "
            "write operations (save / update / delete) using JSON-serialised "
            "Pydantic models. CaseReadRepository handles read-only retrieval of "
            "case data, supporting lazy hydration of case state for the reasoning "
            "workflow."
        ),
    },
    "backend/infra/case_search_client.py": {
        "purpose": "Azure AI Search client for the cases index.",
        "summary": (
            "CaseSearchClient wraps the azure-search-documents SDK and provides "
            "hybrid_search() (keyword + vector fusion), upload_documents(), and "
            "delete_documents() for the cases index. Accepts OData filter "
            "expressions to narrow results by status, country, and case ID."
        ),
    },
    "backend/infra/embeddings.py": {
        "purpose": "Azure OpenAI text-embedding client.",
        "summary": (
            "EmbeddingClient calls the Azure OpenAI text-embedding-3-small (or "
            "configured) deployment to produce float-vector embeddings for "
            "queries and documents. Used by HybridRetriever and all ingestion "
            "services to populate the vector field in Azure AI Search indexes."
        ),
    },
    "backend/infra/evidence_search_client.py": {
        "purpose": "Azure AI Search client for the evidence index.",
        "summary": (
            "EvidenceSearchClient mirrors CaseSearchClient for the evidence index, "
            "supporting hybrid search, document upload, and deletion. Evidence "
            "documents are structured investigation artefacts attached to cases "
            "within the D1–D8 workflow."
        ),
    },
    "backend/infra/knowledge_search_client.py": {
        "purpose": "Azure AI Search client for the knowledge / document library index.",
        "summary": (
            "KnowledgeSearchClient provides hybrid search, upload, and deletion "
            "operations against the knowledge index, which stores chunked "
            "technical reference documents uploaded by operators. Results are "
            "retrieved by the similarity and strategy nodes to ground AI answers."
        ),
    },
    "backend/infra/language_model_client.py": {
        "purpose": "Azure OpenAI chat-completion client.",
        "summary": (
            "LanguageModelClient wraps the openai SDK's AzureOpenAI client and "
            "exposes a chat_completion(messages, model, temperature, …) method "
            "returning the assistant message content. Used by all LangGraph "
            "workflow nodes that require LLM calls."
        ),
    },
    "backend/infra/llm_logging_client.py": {
        "purpose": "Decorator/proxy that adds structured logging around LLM calls.",
        "summary": (
            "LoggedLanguageModelClient wraps LanguageModelClient and logs each "
            "prompt, model name, token usage, and latency using Python's logging "
            "module. Enables observability of all AI calls without changing node "
            "code. Conforms to the same interface as LanguageModelClient."
        ),
    },
    # ── Ingestion ─────────────────────────────────────────────────────────
    "backend/ingestion/case_ingestion.py": {
        "purpose": "Services that ingest, update, and index case records.",
        "summary": (
            "CaseIngestionService orchestrates creating or updating a case: it "
            "serialises the case payload, stores it in Blob Storage via "
            "CaseRepository, generates embeddings, and upserts the document into "
            "the Azure AI Search case index. CaseEntryService validates the "
            "incoming action (CREATE_CASE, UPDATE_CASE, CLOSE_CASE). "
            "CaseSearchIndex encapsulates index-creation logic."
        ),
    },
    "backend/ingestion/evidence_ingestion.py": {
        "purpose": "Services for ingesting binary evidence files into Azure Search.",
        "summary": (
            "EvidenceIngestionService accepts uploaded evidence blobs (images, "
            "PDFs, text), extracts metadata, generates embeddings for text "
            "content, and upserts records into the evidence Azure AI Search index. "
            "EvidenceSearchIndex manages index schema definition and creation."
        ),
    },
    "backend/ingestion/knowledge_ingestion.py": {
        "purpose": "Services for chunking and indexing knowledge documents.",
        "summary": (
            "KnowledgeIngestionService reads uploaded knowledge files, splits "
            "them into overlapping text chunks, generates embeddings for each "
            "chunk, and bulk-upserts them into the knowledge Azure AI Search "
            "index. KnowledgeSearchIndex manages the index schema."
        ),
    },
    # ── Retrieval ─────────────────────────────────────────────────────────
    "backend/retrieval/hybrid_retriever.py": {
        "purpose": "Orchestrates hybrid (keyword + vector) retrieval across all three indexes.",
        "summary": (
            "HybridRetriever exposes retrieve_similar_cases(), "
            "retrieve_evidence(), and retrieve_knowledge() methods. Each method "
            "generates a query embedding, builds an OData filter expression "
            "(filtering by status, case_id exclusion, country), and calls the "
            "appropriate Azure AI Search client. Results are mapped to typed "
            "summary objects (CaseSummary, EvidenceSummary, KnowledgeSummary)."
        ),
    },
    "backend/retrieval/models.py": {
        "purpose": "Pydantic models for retrieval result payloads.",
        "summary": (
            "Defines CaseSummary, EvidenceSummary, and KnowledgeSummary data "
            "classes used to represent normalised retrieval hits. These are the "
            "typed containers passed from HybridRetriever to LangGraph workflow "
            "nodes."
        ),
    },
    # ── State ─────────────────────────────────────────────────────────────
    "backend/state/incident_state.py": {
        "purpose": "Pydantic models representing the runtime state of an incident case.",
        "summary": (
            "IncidentState holds case_id, case_status, organization_country, and "
            "a reasoning_state dict that tracks D1–D8 structured problem-solving "
            "progress. from_payload() parses raw API payloads and extracts country "
            "from nested d_states structures. Used as the primary context object "
            "loaded by the context graph node at the start of each inference run."
        ),
    },
    # ── Tools ─────────────────────────────────────────────────────────────
    "backend/tools/kpi_tool.py": {
        "purpose": "LangChain/LangGraph tool that computes KPI analytics from case data.",
        "summary": (
            "KPIAnalyticsTool is a callable tool node that calculates operational "
            "KPIs (mean-time-to-resolve, recurrence rates, severity distributions) "
            "from a set of retrieved case records. Results populate the kpi_metrics "
            "field in the LangGraph state for downstream KPI interpretation."
        ),
    },
    # ── Workflow ──────────────────────────────────────────────────────────
    "backend/workflow/models.py": {
        "purpose": "Typed payload and result models used across LangGraph workflow nodes.",
        "summary": (
            "Declares Pydantic/TypedDict models for every inter-node payload: "
            "IntentClassificationResult, OperationalPayload, "
            "SimilarityPayload, "
            "StrategyPayload, KPIMetrics, "
            "KPIInterpretation, ReflectionResult, ContextNodeOutput, and "
            "FinalResponsePayload. Provides static typing across the graph state."
        ),
    },
    "backend/workflow/unified_incident_graph.py": {
        "purpose": "Main LangGraph StateGraph definition for the AI reasoning workflow.",
        "summary": (
            "UnifiedIncidentGraph builds the LangGraph StateGraph (IncidentGraphState) "
            "that drives the D1–D8 incident reasoning pipeline. It wires together "
            "14 specialised nodes: start, context, intent classification + "
            "reflection, router, operational + reflection + escalation, similarity "
            "+ reflection, strategy + reflection + escalation, KPI + reflection, "
            "response formatter, and end. Conditional edges use EscalationController "
            "to select the correct next node based on reflection results."
        ),
    },
    # ── Workflow nodes ────────────────────────────────────────────────────
    "backend/workflow/nodes/start_node.py": {
        "purpose": "LangGraph start node — initialises graph state from the entry payload.",
        "summary": (
            "StartNode copies the incoming question and case_id into the graph "
            "state, setting the initial current_d_state. Acts as the entry point "
            "of every reasoning invocation."
        ),
    },
    "backend/workflow/nodes/context_node.py": {
        "purpose": "LangGraph context node — loads full case context from Blob Storage.",
        "summary": (
            "ContextNode reads the case record from CaseReadRepository, extracts "
            "the D-state history and metadata, and populates the case_context "
            "field in state. Subsequent nodes receive the full case context "
            "without needing to query storage directly."
        ),
    },
    "backend/workflow/nodes/intent_classification_node.py": {
        "purpose": "LangGraph node — classifies user query intent using gpt-4o-mini.",
        "summary": (
            "IntentClassificationNode calls the LLM with a structured prompt to "
            "classify the incoming question into one of the routing categories "
            "(operational, similarity, strategy, kpi). Outputs an "
            "IntentClassificationResult stored in state.classification."
        ),
    },
    "backend/workflow/nodes/intent_reflection_node.py": {
        "purpose": "LangGraph reflection node — validates and corrects intent classification.",
        "summary": (
            "IntentReflectionNode re-evaluates the classification result and "
            "checks for ambiguous or incorrect intent assignments. May override "
            "the classification before the router node proceeds."
        ),
    },
    "backend/workflow/nodes/router_node.py": {
        "purpose": "LangGraph router node — sets the route field that drives conditional edges.",
        "summary": (
            "RouterNode reads state.classification and writes state.route to "
            "direct the graph to the operational, similarity, strategy, or KPI "
            "sub-graph branch."
        ),
    },
    "backend/workflow/nodes/operational_node.py": {
        "purpose": "LangGraph node — generates operational guidance using retrieved evidence.",
        "summary": (
            "OperationalNode calls HybridRetriever to fetch relevant evidence "
            "and knowledge documents, then invokes the LLM (default gpt-4o-mini) "
            "to generate step-by-step operational guidance aligned with the "
            "current D-state. Stores results in state.operational_draft."
        ),
    },
    "backend/workflow/nodes/operational_reflection_node.py": {
        "purpose": "LangGraph reflection node — critiques and scores operational guidance.",
        "summary": (
            "OperationalReflectionNode evaluates the operational draft for "
            "completeness, safety, and alignment with D-state goals. Outputs "
            "a ReflectionResult that may set needs_escalation=True to trigger "
            "a premium model retry."
        ),
    },
    "backend/workflow/nodes/operational_escalation_node.py": {
        "purpose": "LangGraph escalation node — retries operational guidance with gpt-4o.",
        "summary": (
            "OperationalEscalationNode re-runs the operational prompt with the "
            "premium model deployment after the reflection node requests "
            "escalation. Sets state.operational_escalated=True to prevent "
            "infinite loops."
        ),
    },
    "backend/workflow/nodes/similarity_node.py": {
        "purpose": "LangGraph node — retrieves and summarises similar historical cases.",
        "summary": (
            "SimilarityNode calls HybridRetriever.retrieve_similar_cases() with "
            "an embedding of the current question, applies country and status "
            "filters, and instructs the LLM to synthesise a comparative analysis "
            "of the top-k matches. Results go into state.similarity_result."
        ),
    },
    "backend/workflow/nodes/similarity_reflection_node.py": {
        "purpose": "LangGraph reflection node — validates the similarity analysis.",
        "summary": (
            "SimilarityReflectionNode checks the similarity result for relevance "
            "and factual grounding. May request a rerun with modified retrieval "
            "parameters if the initial results are deemed insufficient."
        ),
    },
    "backend/workflow/nodes/strategy_node.py": {
        "purpose": "LangGraph node — derives root-cause hypotheses and D8-level strategy.",
        "summary": (
            "StrategyNode aggregates the operational result, similarity matches, "
            "and retrieved knowledge chunks, then calls the LLM to generate a "
            "structured root-cause analysis and corrective action strategy. "
            "Targets D5–D8 stages of the problem-solving workflow."
        ),
    },
    "backend/workflow/nodes/strategy_reflection_node.py": {
        "purpose": "LangGraph reflection node — critiques the strategy recommendation.",
        "summary": (
            "StrategyReflectionNode reviews the strategy draft for logical "
            "consistency, evidence backing, and feasibility. A failing reflection "
            "score triggers escalation to gpt-4o for a higher-quality response."
        ),
    },
    "backend/workflow/nodes/strategy_escalation_node.py": {
        "purpose": "LangGraph escalation node — retries strategy generation with gpt-4o.",
        "summary": (
            "StrategyEscalationNode re-invokes the strategy prompt with the "
            "premium model after reflection-triggered escalation. Sets "
            "state.strategy_escalated=True and stores the improved output."
        ),
    },
    "backend/workflow/nodes/kpi_node.py": {
        "purpose": "LangGraph node — executes KPI analytics computation.",
        "summary": (
            "KPINode invokes KPIAnalyticsTool against the retrieved case set to "
            "compute mean-time-to-resolve, recurrence, and severity trend metrics. "
            "Populates state.kpi_metrics for the KPI reflection and response "
            "formatter nodes."
        ),
    },
    "backend/workflow/nodes/kpi_reflection_node.py": {
        "purpose": "LangGraph reflection node — validates KPI metrics and interpretation.",
        "summary": (
            "KPIReflectionNode checks that computed KPI values are within "
            "plausible ranges and that the interpretation narrative accurately "
            "describes the metrics. May flag anomalies for operator review."
        ),
    },
    "backend/workflow/nodes/response_formatter_node.py": {
        "purpose": "LangGraph node — assembles the final structured JSON response.",
        "summary": (
            "ResponseFormatterNode collects results from whichever branch "
            "(operational / similarity / strategy / KPI) was activated and "
            "formats a FinalResponsePayload with consistent structure including "
            "answer, source references, confidence, and D-state metadata."
        ),
    },
    "backend/workflow/nodes/end_node.py": {
        "purpose": "LangGraph terminal node — finalises graph execution.",
        "summary": (
            "EndNode performs any cleanup, marks the graph run as complete, and "
            "emits the final state. Acts as the sink node for all branches of "
            "the UnifiedIncidentGraph."
        ),
    },
    # ── Tests ─────────────────────────────────────────────────────────────
    "tests/__init__.py": {
        "purpose": "Package initialiser for the tests module.",
        "summary": "Marks the tests directory as a Python package for pytest discovery.",
    },
    "tests/unit/__init__.py": {
        "purpose": "Package initialiser for the unit tests sub-module.",
        "summary": "Marks the unit tests directory as a Python package.",
    },
    "tests/unit/test_escalation_controller.py": {
        "purpose": "Unit tests for EscalationController.",
        "summary": (
            "Covers should_escalate_operational() and should_escalate_strategy() "
            "with parameterised test cases: no reflection present, reflection with "
            "needs_escalation=False, reflection with needs_escalation=True, and "
            "guard against double-escalation via the escalated flag."
        ),
    },
    "tests/unit/test_model_policy.py": {
        "purpose": "Unit tests for ModelPolicy.resolve_model().",
        "summary": (
            "Verifies that the correct deployment name is returned for each "
            "combination of node_name and escalation state. Tests cover "
            "operational default, operational premium, strategy default, "
            "strategy premium, and the fallback intent model."
        ),
    },
    "tests/unit/test_operational_adaptive.py": {
        "purpose": "Integration-style unit tests for the operational adaptive flow.",
        "summary": (
            "Tests the end-to-end path from OperationalNode through "
            "OperationalReflectionNode to OperationalEscalationNode, using mocked "
            "LLM and retriever clients. Validates that escalation is triggered "
            "correctly and that state transitions match expectations."
        ),
    },
    # ── Docs ──────────────────────────────────────────────────────────────
    "docs/case_schema_v1.json": {
        "purpose": "JSON schema definition for the D1–D8 case document structure (v1).",
        "summary": (
            "Defines the canonical schema for case records stored in Azure Blob "
            "Storage and indexed in Azure AI Search. Includes fields for case_id, "
            "organisation, d_states (D1_2 through D8), timestamps, status, "
            "severity, and metadata for each problem-solving stage."
        ),
    },
    # ── UI ────────────────────────────────────────────────────────────────
    "ui/index.html": {
        "purpose": "Single-page HTML shell for the operator decision-support UI.",
        "summary": (
            "The main HTML page that loads the CSS stylesheet and JavaScript "
            "module. Hosts the chat interface, case context panel, and "
            "D1–D8 stage tracker that operators interact with."
        ),
    },
    "ui/styles.css": {
        "purpose": "CSS stylesheet for the operator decision-support UI.",
        "summary": (
            "Defines layout, colour scheme, typography, and component styles "
            "for the chat interface, incident context sidebar, response cards, "
            "and D-state progress indicators."
        ),
    },
    "ui/ui.js": {
        "purpose": "Frontend JavaScript module for the operator decision-support UI.",
        "summary": (
            "Handles user interactions: sending messages to the FastAPI backend "
            "via POST /entry/reasoning, rendering AI responses, displaying "
            "retrieved cases and knowledge sources, and updating the D-state "
            "progress tracker."
        ),
    },
    "ui/staticwebapp.config.json": {
        "purpose": "Azure Static Web Apps routing and security configuration.",
        "summary": (
            "Configures route rules, fallback behaviour, and allowed HTTP methods "
            "for Azure Static Web Apps deployment. Ensures SPA routing works "
            "correctly and may include header/CORS overrides."
        ),
    },
    "ui/commit_commands": {
        "purpose": "Shell script or text file recording git commit commands for UI changes.",
        "summary": "Contains git add / commit / push commands used to deploy UI updates.",
    },
    # ── Artifacts ─────────────────────────────────────────────────────────
    "ARTIFACTS/DECISION-SUPPORT-ARCH.drawio": {
        "purpose": "Architecture diagram (draw.io XML) of the decision-support platform.",
        "summary": "Visual overview of the system architecture.",
    },
    "ARTIFACTS/DECISION-SUPPORT-ARCH-20260213.drawio": {
        "purpose": "Versioned architecture diagram dated 2026-02-13.",
        "summary": "Snapshot of the system architecture as of 13 February 2026.",
    },
    "ARTIFACTS/DECISION-SUPPORT-ARCH-20260213-01.drawio": {
        "purpose": "First revision of the 2026-02-13 architecture diagram.",
        "summary": "Revision 01 of the architecture overview diagram.",
    },
    "ARTIFACTS/DECISION-SUPPORT-ARCH-20260213-02.drawio": {
        "purpose": "Second revision of the 2026-02-13 architecture diagram.",
        "summary": "Revision 02 of the architecture overview diagram.",
    },
}


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------


def should_exclude(path: Path) -> bool:
    for part in path.parts:
        if part in EXCLUDED_DIRS:
            return True
    if path.suffix in EXCLUDED_EXTENSIONS:
        return True
    return False


def get_imports(text: str) -> list[str]:
    """Extract top-level import lines from Python/JS/TS source."""
    imports: list[str] = []
    for line in text.splitlines():
        stripped = line.strip()
        if stripped.startswith(("import ", "from ")):
            imports.append(stripped)
        elif imports and not stripped:
            # Allow one blank line gap in import block
            continue
        elif imports and not stripped.startswith(("import ", "from ", "#")):
            break
    return imports[:30]  # cap at 30 lines


def auto_summary(file_path: Path, content: str) -> str:
    """Generate a simple auto-summary from file content for non-curated files."""
    lines = [l for l in content.splitlines() if l.strip()]
    preview = "\n".join(lines[:8])
    return (
        f"(Auto) File contains {len(lines)} non-empty lines. Preview: {preview[:300]}"
    )


def read_file_safe(path: Path) -> Optional[str]:
    if path.suffix not in SUMMARISABLE_EXTENSIONS:
        return None
    try:
        return path.read_text(encoding="utf-8", errors="replace")[:MAX_READ_CHARS]
    except Exception:
        return None


def collect_files(root: Path) -> list[Path]:
    result = []
    for p in sorted(root.rglob("*")):
        if p.is_file() and not should_exclude(p):
            result.append(p)
    return result


def rel_posix(root: Path, p: Path) -> str:
    return p.relative_to(root).as_posix()


def top_level_folder(root: Path, p: Path) -> str:
    parts = p.relative_to(root).parts
    return parts[0] if len(parts) > 1 else "(root)"


# ---------------------------------------------------------------------------
# Document styling helpers
# ---------------------------------------------------------------------------


def add_toc(doc: Document):
    """Insert a Word field-based Table of Contents."""
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = 'TOC \\o "1-3" \\h \\z \\u'
    fld_char_separate = OxmlElement("w:fldChar")
    fld_char_separate.set(qn("w:fldCharType"), "separate")
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_separate)
    run._r.append(fld_char_end)


def set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def add_field_row(table, label: str, value: str, label_bg="D9E1F2"):
    row = table.add_row()
    lbl_cell = row.cells[0]
    val_cell = row.cells[1]
    set_cell_bg(lbl_cell, label_bg)
    lbl_cell.text = label
    lbl_cell.paragraphs[0].runs[0].bold = True
    val_cell.text = value


# ---------------------------------------------------------------------------
# Main document generator
# ---------------------------------------------------------------------------


def generate_document(root: Path, output: Path):
    files = collect_files(root)

    # Skip the generator script and output file itself
    files = [
        f for f in files if f.name not in {"generate_project_docs.py", output.name}
    ]

    doc = Document()

    # --- Page margins ---
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.2)
        section.right_margin = Inches(1.2)

    # ── Cover Page ────────────────────────────────────────────────────────
    doc.add_heading("valuesims-decision-support", 0)

    cover = doc.add_paragraph()
    cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cover.add_run("Project File Catalogue\n").bold = True
    cover.add_run(f"Generated: {date.today().isoformat()}\n")
    cover.add_run(f"Total files documented: {len(files)}\n")

    top_folders = sorted({top_level_folder(root, f) for f in files})
    cover.add_run(f"Top-level folders: {', '.join(top_folders)}\n")

    doc.add_page_break()

    # ── Table of Contents ─────────────────────────────────────────────────
    doc.add_heading("Table of Contents", 1)
    doc.add_paragraph(
        "Right-click and select 'Update Field' in Microsoft Word to populate the TOC."
    ).italic = True
    add_toc(doc)
    doc.add_page_break()

    # ── Project Overview ──────────────────────────────────────────────────
    doc.add_heading("Project Overview", 1)
    doc.add_paragraph(
        "valuesims-decision-support is an AI-Assisted Incident & Decision-Support "
        "Platform for industrial organisations. It implements a structured D1–D8 "
        "8-Disciplines problem-solving methodology powered by a LangGraph "
        "multi-agent reasoning workflow, FastAPI backend, Azure AI Search hybrid "
        "retrieval, and Azure OpenAI (adaptive gpt-4o-mini / gpt-4o model policy)."
    )
    doc.add_paragraph(
        "Key architectural components: FastAPI application factory (app.py), "
        "UnifiedIncidentGraph (LangGraph StateGraph), HybridRetriever (Azure AI "
        "Search), EscalationController + ModelPolicy (adaptive model selection), "
        "ingestion services (case / evidence / knowledge), and a static web UI "
        "deployed on Azure Static Web Apps."
    )

    # ── Sections per top-level folder ─────────────────────────────────────
    grouped: dict[str, list[Path]] = {}
    for f in files:
        folder = top_level_folder(root, f)
        grouped.setdefault(folder, []).append(f)

    for folder in sorted(grouped.keys()):
        doc.add_heading(f"/{folder}" if folder != "(root)" else "/ (Project Root)", 1)

        folder_files = sorted(grouped[folder])

        for file_path in folder_files:
            rel = rel_posix(root, file_path)
            ext = file_path.suffix or "(none)"
            curated = CURATED.get(rel, {})
            content = read_file_safe(file_path)

            purpose = curated.get("purpose", f"{file_path.name} — {ext} file.")
            summary = curated.get(
                "summary",
                (
                    auto_summary(file_path, content)
                    if content
                    else "(Binary or unreadable file.)"
                ),
            )

            imports: list[str] = []
            if content and file_path.suffix in {".py", ".ts", ".js"}:
                imports = get_imports(content)

            # Sub-heading: filename
            doc.add_heading(file_path.name, 3)

            # Details table
            table = doc.add_table(rows=0, cols=2)
            table.style = "Table Grid"
            table.columns[0].width = Inches(1.6)
            table.columns[1].width = Inches(4.8)

            add_field_row(table, "Relative Path", rel)
            add_field_row(table, "File Type", ext)
            add_field_row(table, "Purpose", purpose)
            add_field_row(table, "Summary", summary)

            if imports:
                add_field_row(table, "Key Imports", "\n".join(imports[:20]))

            doc.add_paragraph()  # spacing

    # ── Appendix ──────────────────────────────────────────────────────────
    doc.add_page_break()
    doc.add_heading("Appendix: All File Paths", 1)
    doc.add_paragraph("Flat list of all documented file paths for quick reference.")
    doc.add_paragraph()

    for f in files:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(rel_posix(root, f)).font.size = Pt(9)

    # ── Save ──────────────────────────────────────────────────────────────
    doc.save(str(output))
    return len(files), sorted(top_folders)


# ---------------------------------------------------------------------------
# Entry point
# ---------------------------------------------------------------------------

if __name__ == "__main__":
    print(f"Project root : {PROJECT_ROOT}")
    print(f"Output file  : {OUTPUT_FILE}")
    print("Walking file tree …")

    count, folders = generate_document(PROJECT_ROOT, OUTPUT_FILE)

    print(f"\n✓ Document created: {OUTPUT_FILE}")
    print(f"  Files documented : {count}")
    print(f"  Top-level folders: {', '.join(folders)}")
    size_kb = OUTPUT_FILE.stat().st_size / 1024
    print(f"  File size        : {size_kb:.1f} KB")
